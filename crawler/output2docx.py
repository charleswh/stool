# coding=utf-8
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from crawler.blakfp import PAPER_DIR

def add_text(para, text, bold=False, color=None):
    if color != None and color[0:1] != '#':
        if color in ('red', 'Red'):
            color = '#ff0000'
    run = para.add_run(text)
    run.font.name = u'微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    if bold != False:
        run.bold = True
    if color != None:
        r = int(color[1:3], 16)
        g = int(color[3:5], 16)
        b = int(color[5:7], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def add_pic(para, pic_link, opener=None):
    import urllib
    request = urllib.request.Request(pic_link)
    try:
        response = opener.open(request)
    except urllib.error.HTTPError as err:
        if err.code == 404:
            return
    pic_content = response.read()
    local_tmp_pic = pic_link.split('/')[-1]
    with open(local_tmp_pic, 'wb') as f:
        f.write(pic_content)
    para.add_run().add_picture(local_tmp_pic, width=Cm(15))
    os.unlink(local_tmp_pic)


def output2docx(text=None, docx_name=None, opener=None):
    if text == None:
        pass
        print('Need text to output!')
        exit(0)
    if not isinstance(text, list):
        print('Text format error, must be LIST!')
        exit(0)
    if docx_name == None:
        docx_name = 'demo.docx'

    print('Write paper content to docx file %s ...' % docx_name)
    inter_content = list(text)
    title = inter_content.pop(0)
    docx = Document()
    docx_title = docx.add_heading()
    docx_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docx_title.paragraph_format.space_after = Pt(20)
    title_run = docx_title.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.name = u'微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    docx_para = docx.add_paragraph()
    for top in inter_content:
        if isinstance(top, str):
            if top == '\n':
                docx_para = docx.add_paragraph()
            else:
                add_text(docx_para, top)
        elif len(top) == 2 and top[1] == 'bold':
            add_text(docx_para, top[0], True)
        elif len(top) == 2 and (top[1][0:1] == '#' or top[1] in ('red', 'Red')):
            add_text(docx_para, top[0], color=top[1])
        elif len(top) == 2 and top[1] == 'picture':
            add_pic(docx_para, top[0], opener)
        else:
            for sub in top:
                if isinstance(sub, str):
                    if sub == '\n':
                        docx_para = docx.add_paragraph()
                    else:
                        add_text(docx_para, sub)
                elif len(sub) == 2 and (sub[1][0:1] == '#' or sub[1] in ('red', 'Red')):
                    add_text(docx_para, sub[0], color=sub[1])
                else:
                    for bot in sub:
                        if bot == '\n':
                            docx_para = docx.add_paragraph()
                        else:
                            add_text(docx_para, bot)

    if not os.path.exists(PAPER_DIR):
        os.makedirs(PAPER_DIR)
    docx.save(os.path.join(PAPER_DIR, docx_name))
    print('Write paper done.')


if __name__ == '__main__':
    output2docx()
